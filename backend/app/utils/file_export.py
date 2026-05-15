"""
File export utilities – generate structured, readable PDF and DOCX brand kits.
All sections are rendered as human-readable text, not raw JSON.
"""
from pathlib import Path
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor, Color, white, black
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    Table, TableStyle, HRFlowable,
)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _str(val) -> str:
    if val is None:
        return ""
    if isinstance(val, list):
        return ", ".join(str(v) for v in val if v)
    if isinstance(val, dict):
        return "; ".join(f"{k}: {v}" for k, v in val.items() if v)
    return str(val).strip()


def _list(val) -> list:
    if isinstance(val, list):
        return [str(v) for v in val if v]
    if val:
        return [str(val)]
    return []


def _hex(val: str) -> str:
    """Return a valid 6-digit hex string, defaulting to slate if invalid."""
    v = str(val or "").strip().lstrip("#")
    return f"#{v}" if len(v) == 6 else "#334155"


# ── PDF ───────────────────────────────────────────────────────────────────────

_DARK   = HexColor("#0f172a")
_MID    = HexColor("#1e293b")
_RULE   = HexColor("#334155")
_BODY   = HexColor("#e2e8f0")
_MUTED  = HexColor("#94a3b8")
_WHITE  = HexColor("#ffffff")
_ACCENT = HexColor("#6366f1")


def _pdf_styles():
    base = getSampleStyleSheet()

    def add(name, **kw):
        try:
            base.add(ParagraphStyle(name=name, **kw))
        except KeyError:
            pass  # already registered

    add("Cover",        fontSize=32, leading=38, textColor=_WHITE,   spaceAfter=8,  alignment=1, fontName="Helvetica-Bold")
    add("CoverSub",     fontSize=15, leading=20, textColor=_MUTED,   spaceAfter=4,  alignment=1)
    add("SectionTitle", fontSize=14, leading=18, textColor=_ACCENT,  spaceAfter=6,  spaceBefore=18, fontName="Helvetica-Bold")
    add("FieldLabel",   fontSize=8,  leading=11, textColor=_MUTED,   spaceAfter=2,  fontName="Helvetica-Bold", wordWrap="CJK")
    add("FieldValue",   fontSize=10, leading=15, textColor=_BODY,    spaceAfter=8)
    add("Bullet",       fontSize=10, leading=14, textColor=_BODY,    spaceAfter=4,  leftIndent=14)
    add("Quote",        fontSize=11, leading=16, textColor=_WHITE,   spaceAfter=8,  leftIndent=16, fontName="Helvetica-Oblique")
    add("Tag",          fontSize=9,  leading=13, textColor=_ACCENT,  spaceAfter=4)
    return base


def _rule(elements):
    elements.append(HRFlowable(width="100%", thickness=0.5, color=_RULE, spaceAfter=10))


def _label(elements, text, styles):
    elements.append(Paragraph(text.upper(), styles["FieldLabel"]))


def _value(elements, text, styles):
    if text:
        elements.append(Paragraph(str(text), styles["FieldValue"]))


def _bullets(elements, items, styles):
    for item in _list(items):
        elements.append(Paragraph(f"• {item}", styles["Bullet"]))


def _color_chip(hex_val: str, label: str, styles) -> Table:
    """A small inline color chip table for the PDF."""
    chip_color = HexColor(_hex(hex_val).lstrip("#") and _hex(hex_val))
    try:
        chip_color = HexColor(_hex(hex_val))
    except Exception:
        chip_color = HexColor("#334155")

    t = Table([[" ", f"{label}  {hex_val}"]],
              colWidths=[14, 200], rowHeights=[14])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), chip_color),
        ("TEXTCOLOR",  (1, 0), (1, 0), _BODY),
        ("FONTSIZE",   (1, 0), (1, 0), 9),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING",    (0, 0), (-1, -1), 2),
    ]))
    return t


def _section(elements, title, styles):
    elements.append(Spacer(1, 14))
    elements.append(Paragraph(title, styles["SectionTitle"]))
    _rule(elements)


def generate_pdf(project_id: str, brand_data: dict) -> str:
    filepath = EXPORT_DIR / f"{project_id}_brand_kit.pdf"

    doc = SimpleDocTemplate(
        str(filepath), pagesize=A4,
        topMargin=0.65 * inch, bottomMargin=0.65 * inch,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
    )
    styles = _pdf_styles()
    el = []

    naming  = brand_data.get("naming",         {}) or {}
    strat   = brand_data.get("brand_strategy", {}) or {}
    content = brand_data.get("content_agent",  {}) or {}
    guide   = brand_data.get("guidelines_agent", {}) or {}
    visual  = brand_data.get("visual_identity_agent", {}) or {}
    idea    = brand_data.get("idea_discovery", {}) or {}
    market  = brand_data.get("market_research", {}) or {}
    comp    = brand_data.get("competitor_analysis", {}) or {}

    brand_name = _str(naming.get("brand_name") or "Brand Kit")
    tagline    = _str(naming.get("tagline", ""))
    mission    = _str(content.get("mission_statement") or strat.get("brand_mission", ""))
    vision     = _str(content.get("vision_statement")  or strat.get("brand_vision",  ""))

    # ── Cover ──────────────────────────────────────────────────────────────
    el.append(Spacer(1, 1.8 * inch))
    el.append(Paragraph(brand_name, styles["Cover"]))
    if tagline:
        el.append(Paragraph(tagline, styles["CoverSub"]))
    el.append(Spacer(1, 0.3 * inch))
    el.append(Paragraph("Brand Identity Kit", styles["CoverSub"]))
    el.append(Paragraph(datetime.now().strftime("Generated %B %Y"), styles["FieldLabel"]))
    el.append(PageBreak())

    # ── 1. Brand Foundation ────────────────────────────────────────────────
    _section(el, "1. Brand Foundation", styles)
    if mission:
        _label(el, "Mission", styles)
        el.append(Paragraph(f'"{mission}"', styles["Quote"]))
    if vision:
        _label(el, "Vision", styles)
        el.append(Paragraph(f'"{vision}"', styles["Quote"]))
    stands = _str(content.get("brand_stands_for", ""))
    if stands:
        _label(el, "What We Stand For", styles)
        _value(el, stands, styles)
    values = _list(strat.get("brand_values") or guide.get("brand_overview", {}).get("values", []))
    if values:
        _label(el, "Core Values", styles)
        _bullets(el, values, styles)

    # ── 2. Brand Strategy ──────────────────────────────────────────────────
    _section(el, "2. Brand Strategy", styles)
    bp = strat.get("brand_personality", {}) or {}
    for lbl, key in [("Archetype", "archetype"), ("Tone of Voice", "tone_of_voice"), ("Brand Promise", "brand_promise")]:
        v = _str(bp.get(key) or strat.get(key, ""))
        if v:
            _label(el, lbl, styles); _value(el, v, styles)
    pos = _str(strat.get("positioning_statement", ""))
    if pos:
        _label(el, "Positioning Statement", styles); _value(el, pos, styles)
    segs = _list(strat.get("target_segments", []))
    if segs:
        _label(el, "Target Audience", styles); _bullets(el, segs, styles)

    # ── 3. Brand Naming ────────────────────────────────────────────────────
    _section(el, "3. Brand Naming", styles)
    for lbl, key in [("Brand Name", "brand_name"), ("Tagline", "tagline"), ("Name Rationale", "name_rationale"), ("Tagline Rationale", "tagline_rationale")]:
        v = _str(naming.get(key, ""))
        if v:
            _label(el, lbl, styles); _value(el, v, styles)
    alts = _list(naming.get("alternative_names", []))
    if alts:
        _label(el, "Alternative Names Considered", styles); _bullets(el, alts, styles)

    # ── 4. Visual Identity ─────────────────────────────────────────────────
    _section(el, "4. Visual Identity", styles)
    primary = _str(visual.get("primary_color", ""))
    accent  = _str(visual.get("accent_color",  ""))
    font    = _str(visual.get("font", ""))
    if primary:
        _label(el, "Primary Colour", styles)
        el.append(_color_chip(primary, "Primary", styles))
        el.append(Spacer(1, 4))
    if accent:
        _label(el, "Accent Colour", styles)
        el.append(_color_chip(accent, "Accent", styles))
        el.append(Spacer(1, 4))
    if font:
        _label(el, "Primary Typeface", styles); _value(el, font, styles)

    # Color rationale
    cr = guide.get("color_rationale", {}) or {}
    if cr.get("primary_reasoning"):
        _label(el, "Why This Primary Colour", styles); _value(el, cr["primary_reasoning"], styles)
    if cr.get("accent_reasoning"):
        _label(el, "Why This Accent Colour", styles);  _value(el, cr["accent_reasoning"], styles)
    if cr.get("palette_harmony"):
        _label(el, "Palette Harmony", styles); _value(el, cr["palette_harmony"], styles)

    # Typography rationale
    tr = guide.get("typography_rationale", {}) or {}
    if tr.get("heading_font_reason"):
        _label(el, "Why This Heading Font", styles); _value(el, tr["heading_font_reason"], styles)
    if tr.get("body_font_reason"):
        _label(el, "Why This Body Font", styles); _value(el, tr["body_font_reason"], styles)

    # Logo concepts (text only)
    concepts = visual.get("design_concepts", []) or []
    if concepts:
        _label(el, "Logo Concepts Overview", styles)
        for c in concepts[:10]:
            name = _str(c.get("name", ""))
            vc   = _str(c.get("visual_concept", ""))
            if name:
                el.append(Paragraph(f"<b>{c.get('number', '')}. {name}</b>", styles["FieldValue"]))
            if vc:
                el.append(Paragraph(vc, styles["Bullet"]))
                el.append(Spacer(1, 4))

    # ── 5. Brand Content ───────────────────────────────────────────────────
    _section(el, "5. Brand Content", styles)
    for lbl, key in [("Elevator Pitch", "elevator_pitch"), ("About Section", "about_section")]:
        v = _str(content.get(key, ""))
        if v:
            _label(el, lbl, styles); _value(el, v, styles)

    email_tag = _str(content.get("email_tagline") or content.get("email_signature_tagline", ""))
    if email_tag:
        _label(el, "Email Signature Tagline", styles)
        el.append(Paragraph(f'"{email_tag}"', styles["Quote"]))

    tov = content.get("tone_of_voice", {}) or {}
    chars = _list(tov.get("character", []))
    if chars:
        _label(el, "Tone of Voice — Character", styles)
        el.append(Paragraph(" · ".join(chars), styles["Tag"]))
    if tov.get("description"):
        _value(el, tov["description"], styles)
    write_like = _list(tov.get("write_like", []))
    if write_like:
        _label(el, "Write Like This", styles)
        _bullets(el, [f'"{w}"' for w in write_like], styles)
    avoid = _list(tov.get("avoid", []))
    if avoid:
        _label(el, "Avoid", styles)
        _bullets(el, avoid, styles)
    if tov.get("example_hero_copy"):
        _label(el, "Sample Brand Voice", styles)
        el.append(Paragraph(f'"{tov["example_hero_copy"]}"', styles["Quote"]))

    # ── 6. Messaging Pillars ───────────────────────────────────────────────
    pillars = content.get("key_messaging_pillars", []) or []
    if pillars:
        _section(el, "6. Messaging Pillars", styles)
        for i, p in enumerate(pillars, 1):
            el.append(Paragraph(f"<b>Pillar {i}: {_str(p.get('pillar', ''))}</b>", styles["FieldValue"]))
            if p.get("headline"):
                el.append(Paragraph(_str(p["headline"]), styles["Quote"]))
            if p.get("description"):
                _value(el, p["description"], styles)
            el.append(Spacer(1, 6))

    # ── 7. Social Media ────────────────────────────────────────────────────
    bios = content.get("social_media_bios", {}) or {}
    if any(bios.values()):
        _section(el, "7. Social Media Bios", styles)
        for platform in ["twitter", "instagram", "linkedin"]:
            bio = _str(bios.get(platform, ""))
            if bio:
                _label(el, platform.title(), styles); _value(el, bio, styles)
    tags = _list(content.get("brand_hashtags", []))
    if tags:
        _label(el, "Brand Hashtags", styles)
        el.append(Paragraph("  ".join(tags), styles["Tag"]))

    ctas = _list(content.get("call_to_action_phrases", []))
    if ctas:
        _label(el, "Call-to-Action Phrases", styles); _bullets(el, ctas, styles)

    # ── 8. Brand Rules ─────────────────────────────────────────────────────
    rules = guide.get("brand_rules", []) or []
    if rules:
        _section(el, "8. Brand Rules", styles)
        for r in rules:
            num  = r.get("rule_number", "")
            name = _str(r.get("rule", ""))
            desc = _str(r.get("description", ""))
            why  = _str(r.get("why", ""))
            el.append(Paragraph(f"<b>Rule {num}: {name}</b>", styles["FieldValue"]))
            if desc: _value(el, desc, styles)
            if why:
                el.append(Paragraph(f"Why: {why}", styles["Bullet"]))
            el.append(Spacer(1, 6))

    # ── 9. Voice & Tone Guidelines ─────────────────────────────────────────
    vt = guide.get("voice_and_tone", {}) or {}
    dos   = _list(vt.get("dos",   []))
    donts = _list(vt.get("donts", []))
    if dos or donts:
        _section(el, "9. Voice & Tone Guidelines", styles)
        if vt.get("personality"):
            _value(el, vt["personality"], styles)
        if dos:
            _label(el, "Do", styles); _bullets(el, dos, styles)
        if donts:
            _label(el, "Don't", styles); _bullets(el, donts, styles)
        for phrase in _list(vt.get("example_phrases", [])):
            el.append(Paragraph(f'"{phrase}"', styles["Quote"]))

    # ── 10. Logo Usage ─────────────────────────────────────────────────────
    logo = guide.get("logo_usage", {}) or {}
    if logo:
        _section(el, "10. Logo Usage Rules", styles)
        if logo.get("primary_usage"):  _label(el, "Usage",      styles); _value(el, logo["primary_usage"], styles)
        if logo.get("minimum_size"):   _label(el, "Min Size",   styles); _value(el, logo["minimum_size"],  styles)
        if logo.get("clear_space"):    _label(el, "Clear Space",styles); _value(el, logo["clear_space"],   styles)
        logo_donts = _list(logo.get("dont_rules", []))
        if logo_donts:
            _label(el, "Don't Rules", styles); _bullets(el, logo_donts, styles)

    # ── 11. Market & Competitors (summary) ────────────────────────────────
    trends = _list(market.get("market_trends", []))
    direct = [c.get("name", "") for c in (comp.get("direct_competitors", []) or [])[:6] if isinstance(c, dict)]
    if trends or direct:
        _section(el, "11. Market & Competitors", styles)
        if direct:
            _label(el, "Key Competitors", styles); _bullets(el, direct, styles)
        if trends:
            _label(el, "Market Trends", styles); _bullets(el, trends[:5], styles)

    doc.build(el)
    return str(filepath)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _set_heading_color(paragraph, hex_color: str):
    """Apply a hex color to a paragraph's run text."""
    r, g, b = (int(hex_color.lstrip("#")[i:i+2], 16) for i in (0, 2, 4))
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _docx_heading(doc, text: str, level: int = 1, color: str | None = None):
    h = doc.add_heading(text, level=level)
    if color:
        _set_heading_color(h, color)
    return h


def _docx_label(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    p.paragraph_format.space_after = Pt(2)
    return p


def _docx_value(doc, text: str):
    if not text:
        return
    p = doc.add_paragraph(str(text))
    p.paragraph_format.space_after = Pt(6)
    return p


def _docx_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(str(text))
    p.paragraph_format.space_after = Pt(3)
    return p


def _docx_quote(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f'"{text}"')
    run.font.italic = True
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    return p


def _docx_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "334155")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(project_id: str, brand_data: dict) -> str:
    filepath = EXPORT_DIR / f"{project_id}_brand_kit.docx"

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    naming  = brand_data.get("naming",               {}) or {}
    strat   = brand_data.get("brand_strategy",       {}) or {}
    content = brand_data.get("content_agent",        {}) or {}
    guide   = brand_data.get("guidelines_agent",     {}) or {}
    visual  = brand_data.get("visual_identity_agent",{}) or {}
    market  = brand_data.get("market_research",      {}) or {}
    comp    = brand_data.get("competitor_analysis",  {}) or {}

    brand_name = _str(naming.get("brand_name") or "Brand Kit")
    tagline    = _str(naming.get("tagline", ""))
    mission    = _str(content.get("mission_statement") or strat.get("brand_mission", ""))
    vision     = _str(content.get("vision_statement")  or strat.get("brand_vision",  ""))

    # ── Cover ──────────────────────────────────────────────────────────────
    title = doc.add_heading(brand_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, "6366f1")

    if tagline:
        sub = doc.add_paragraph(tagline)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size  = Pt(14)
        sub.runs[0].font.italic = True
        sub.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    date_p = doc.add_paragraph(f"Brand Identity Kit  ·  {datetime.now().strftime('%B %Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(9)
    date_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    doc.add_page_break()

    # ── 1. Brand Foundation ────────────────────────────────────────────────
    _docx_heading(doc, "1. Brand Foundation", 1, "6366f1"); _docx_rule(doc)
    if mission:
        _docx_label(doc, "Mission")
        _docx_quote(doc, mission)
    if vision:
        _docx_label(doc, "Vision")
        _docx_quote(doc, vision)
    stands = _str(content.get("brand_stands_for", ""))
    if stands:
        _docx_label(doc, "What We Stand For")
        _docx_value(doc, stands)
    values = _list(strat.get("brand_values") or (guide.get("brand_overview") or {}).get("values", []))
    if values:
        _docx_label(doc, "Core Values")
        for v in values: _docx_bullet(doc, v)

    # ── 2. Brand Strategy ──────────────────────────────────────────────────
    _docx_heading(doc, "2. Brand Strategy", 1, "6366f1"); _docx_rule(doc)
    bp = strat.get("brand_personality", {}) or {}
    for lbl, key in [("Archetype", "archetype"), ("Tone of Voice", "tone_of_voice"), ("Brand Promise", "brand_promise")]:
        v = _str(bp.get(key) or strat.get(key, ""))
        if v: _docx_label(doc, lbl); _docx_value(doc, v)
    pos = _str(strat.get("positioning_statement", ""))
    if pos: _docx_label(doc, "Positioning Statement"); _docx_value(doc, pos)
    segs = _list(strat.get("target_segments", []))
    if segs:
        _docx_label(doc, "Target Audience")
        for s in segs: _docx_bullet(doc, s)

    # ── 3. Brand Naming ────────────────────────────────────────────────────
    _docx_heading(doc, "3. Brand Naming", 1, "6366f1"); _docx_rule(doc)
    for lbl, key in [("Brand Name", "brand_name"), ("Tagline", "tagline"), ("Name Rationale", "name_rationale"), ("Tagline Rationale", "tagline_rationale")]:
        v = _str(naming.get(key, ""))
        if v: _docx_label(doc, lbl); _docx_value(doc, v)
    alts = _list(naming.get("alternative_names", []))
    if alts:
        _docx_label(doc, "Alternative Names")
        for a in alts: _docx_bullet(doc, a)

    # ── 4. Visual Identity ─────────────────────────────────────────────────
    _docx_heading(doc, "4. Visual Identity", 1, "6366f1"); _docx_rule(doc)
    primary = _str(visual.get("primary_color", ""))
    accent  = _str(visual.get("accent_color",  ""))
    font    = _str(visual.get("font", ""))
    if primary: _docx_label(doc, "Primary Colour"); _docx_value(doc, primary)
    if accent:  _docx_label(doc, "Accent Colour");  _docx_value(doc, accent)
    if font:    _docx_label(doc, "Primary Typeface"); _docx_value(doc, font)

    cr = guide.get("color_rationale", {}) or {}
    if cr.get("primary_reasoning"): _docx_label(doc, "Why Primary Colour"); _docx_value(doc, cr["primary_reasoning"])
    if cr.get("accent_reasoning"):  _docx_label(doc, "Why Accent Colour");  _docx_value(doc, cr["accent_reasoning"])
    if cr.get("palette_harmony"):   _docx_label(doc, "Palette Harmony");    _docx_value(doc, cr["palette_harmony"])

    tr = guide.get("typography_rationale", {}) or {}
    if tr.get("heading_font_reason"): _docx_label(doc, "Why Heading Font"); _docx_value(doc, tr["heading_font_reason"])
    if tr.get("body_font_reason"):    _docx_label(doc, "Why Body Font");    _docx_value(doc, tr["body_font_reason"])
    if tr.get("combination_logic"):   _docx_label(doc, "Why They Pair Well"); _docx_value(doc, tr["combination_logic"])

    concepts = visual.get("design_concepts", []) or []
    if concepts:
        _docx_label(doc, "Logo Concepts")
        for c in concepts[:10]:
            name = _str(c.get("name", ""))
            vc   = _str(c.get("visual_concept", ""))
            if name:
                p = doc.add_paragraph()
                run = p.add_run(f"{c.get('number', '')}. {name}")
                run.font.bold = True
            if vc:
                _docx_bullet(doc, vc)

    # ── 5. Brand Content ───────────────────────────────────────────────────
    _docx_heading(doc, "5. Brand Content", 1, "6366f1"); _docx_rule(doc)
    for lbl, key in [("Elevator Pitch", "elevator_pitch"), ("About Section", "about_section")]:
        v = _str(content.get(key, ""))
        if v: _docx_label(doc, lbl); _docx_value(doc, v)

    email_tag = _str(content.get("email_tagline") or content.get("email_signature_tagline", ""))
    if email_tag:
        _docx_label(doc, "Email Signature Tagline")
        _docx_quote(doc, email_tag)

    tov = content.get("tone_of_voice", {}) or {}
    chars = _list(tov.get("character", []))
    if chars:
        _docx_label(doc, "Tone of Voice")
        _docx_value(doc, " · ".join(chars))
    if tov.get("description"):
        _docx_value(doc, tov["description"])
    write_like = _list(tov.get("write_like", []))
    if write_like:
        _docx_label(doc, "Write Like This")
        for w in write_like: _docx_bullet(doc, f'"{w}"')
    avoid = _list(tov.get("avoid", []))
    if avoid:
        _docx_label(doc, "Avoid")
        for a in avoid: _docx_bullet(doc, a)
    if tov.get("example_hero_copy"):
        _docx_label(doc, "Sample Brand Voice")
        _docx_quote(doc, tov["example_hero_copy"])

    # ── 6. Messaging Pillars ───────────────────────────────────────────────
    pillars = content.get("key_messaging_pillars", []) or []
    if pillars:
        _docx_heading(doc, "6. Messaging Pillars", 1, "6366f1"); _docx_rule(doc)
        for i, p in enumerate(pillars, 1):
            h = doc.add_paragraph()
            h.add_run(f"Pillar {i}: {_str(p.get('pillar', ''))}").bold = True
            if p.get("headline"):  _docx_quote(doc, _str(p["headline"]))
            if p.get("description"): _docx_value(doc, _str(p["description"]))

    # ── 7. Social Media ────────────────────────────────────────────────────
    bios = content.get("social_media_bios", {}) or {}
    if any(bios.values()):
        _docx_heading(doc, "7. Social Media Bios", 1, "6366f1"); _docx_rule(doc)
        for platform in ["twitter", "instagram", "linkedin"]:
            bio = _str(bios.get(platform, ""))
            if bio: _docx_label(doc, platform.title()); _docx_value(doc, bio)
    tags = _list(content.get("brand_hashtags", []))
    if tags:
        _docx_label(doc, "Brand Hashtags")
        _docx_value(doc, "  ".join(tags))
    ctas = _list(content.get("call_to_action_phrases", []))
    if ctas:
        _docx_label(doc, "Call-to-Action Phrases")
        for cta in ctas: _docx_bullet(doc, cta)

    # ── 8. Brand Rules ─────────────────────────────────────────────────────
    rules = guide.get("brand_rules", []) or []
    if rules:
        _docx_heading(doc, "8. Brand Rules", 1, "6366f1"); _docx_rule(doc)
        for r in rules:
            p = doc.add_paragraph()
            p.add_run(f"Rule {r.get('rule_number', '')}: {_str(r.get('rule', ''))}").bold = True
            if r.get("description"): _docx_value(doc, _str(r["description"]))
            if r.get("why"):
                pw = doc.add_paragraph()
                pw.add_run("Why: ").bold = True
                pw.add_run(_str(r["why"]))
                pw.paragraph_format.space_after = Pt(6)

    # ── 9. Voice & Tone Guidelines ─────────────────────────────────────────
    vt = guide.get("voice_and_tone", {}) or {}
    if vt:
        _docx_heading(doc, "9. Voice & Tone Guidelines", 1, "6366f1"); _docx_rule(doc)
        if vt.get("personality"): _docx_value(doc, vt["personality"])
        dos   = _list(vt.get("dos",   []))
        donts = _list(vt.get("donts", []))
        if dos:   _docx_label(doc, "Do");    [_docx_bullet(doc, d) for d in dos]
        if donts: _docx_label(doc, "Don't"); [_docx_bullet(doc, d) for d in donts]
        for ph in _list(vt.get("example_phrases", [])):
            _docx_quote(doc, ph)

    # ── 10. Logo Usage ─────────────────────────────────────────────────────
    logo = guide.get("logo_usage", {}) or {}
    if logo:
        _docx_heading(doc, "10. Logo Usage Rules", 1, "6366f1"); _docx_rule(doc)
        if logo.get("primary_usage"): _docx_label(doc, "Usage");       _docx_value(doc, logo["primary_usage"])
        if logo.get("minimum_size"):  _docx_label(doc, "Min Size");    _docx_value(doc, logo["minimum_size"])
        if logo.get("clear_space"):   _docx_label(doc, "Clear Space"); _docx_value(doc, logo["clear_space"])
        for d in _list(logo.get("dont_rules", [])): _docx_bullet(doc, d)

    # ── 11. Market & Competitors ───────────────────────────────────────────
    direct = [c.get("name", "") for c in (comp.get("direct_competitors", []) or [])[:6] if isinstance(c, dict)]
    trends = _list(market.get("market_trends", []))
    if direct or trends:
        _docx_heading(doc, "11. Market & Competitors", 1, "6366f1"); _docx_rule(doc)
        if direct: _docx_label(doc, "Key Competitors"); [_docx_bullet(doc, d) for d in direct]
        if trends: _docx_label(doc, "Market Trends");   [_docx_bullet(doc, t) for t in trends[:5]]

    doc.save(str(filepath))
    return str(filepath)
